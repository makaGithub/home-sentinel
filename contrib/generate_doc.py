#!/usr/bin/env python3
"""
Генерация документа проекта home-sentinel из project.md

Использование:
    python3 contrib/generate_doc.py

Скрипт читает docs/project.md и генерирует docs/project.docx с заданным форматированием.

Форматирование:
- Шрифт: Times New Roman, 12pt
- Межстрочный интервал: 1.15
- Выравнивание: по ширине (для основного текста)
- Поля: слева 2 см, справа 1 см, сверху и снизу 2 см

Поддерживаемые элементы Markdown:
- Заголовки: # ## ###
- Параграфы: обычный текст
- Списки: - или *
- Нумерованные списки: 1. 2. 3.
- Вложенные списки: отступы 4 пробела
- Жирный текст: **text**
- Разрывы страниц: ---
- Оглавление: <!-- TOC --> (автоматически генерируется из заголовков)

Примечание: Оглавление создается как текстовый список заголовков, что обеспечивает
совместимость с Microsoft Word и LibreOffice. Оглавление автоматически обновляется
при изменении заголовков в project.md и повторной генерации документа.
"""

import os
import re
import sys
import argparse

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    print("✅ Все необходимые модули импортированы успешно")
except ImportError as e:
    print(f"❌ Ошибка импорта модулей: {e}")
    print("   Установите python-docx: pip install python-docx")
    sys.exit(1)

# Импортируем вспомогательный модуль для конфигурации
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    from config_helper import (
        load_env_file, get_project_root, generate_output_filename,
        get_build_dir, publish_to_gh_pages, create_github_release
    )
except ImportError as e:
    print(f"⚠️  Предупреждение: не удалось импортировать config_helper: {e}")
    print("   Используются значения по умолчанию")

# Константы форматирования
FONT_NAME = 'Times New Roman'
FONT_SIZE = 12
LINE_SPACING = 1.15

# Размеры заголовков
HEADING_SIZES = {1: 16, 2: 14, 3: 13}

def set_paragraph_formatting(para, font_name=FONT_NAME, font_size=FONT_SIZE, 
                             line_spacing=LINE_SPACING, alignment='justify', bold=False):
    """Устанавливает форматирование параграфа"""
    # Шрифт и размер
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if bold:
            run.bold = True
    
    # Если нет runs, создаем один
    if not para.runs:
        run = para.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if bold:
            run.bold = True
    
    # Выравнивание
    if alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Межстрочный интервал
    para_format = para.paragraph_format
    para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para_format.line_spacing = line_spacing

def add_formatted_paragraph(doc, text, font_size=FONT_SIZE, alignment='justify', bold=False):
    """Добавляет параграф с форматированием"""
    para = doc.add_paragraph()
    
    # Обработка жирного текста **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Жирный текст
            run = para.add_run(part[2:-2])
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
            run.bold = True
        elif part.strip():
            # Обычный текст
            run = para.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
    
    set_paragraph_formatting(para, font_size=font_size, alignment=alignment, bold=bold)
    return para

def add_formatted_heading(doc, text, level, font_size=None):
    """Добавляет заголовок с форматированием"""
    size = font_size or HEADING_SIZES.get(level, FONT_SIZE)
    heading = doc.add_heading(text, level)
    set_paragraph_formatting(heading, font_size=size, alignment='left', bold=True)
    return heading

def add_page_number_footer(section):
    """Добавляет нумерацию страниц в нижний колонтитул"""
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Создаем поле PAGE для номера страницы
    run = para.add_run()
    r_element = run._element
    
    # Начало поля
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r_element.append(fldChar_begin)
    
    # Инструкция PAGE
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    r_element.append(instrText)
    
    # Разделитель
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    r_element.append(fldChar_separate)
    
    # Временный текст
    run_text = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run_text.append(rPr)
    t = OxmlElement('w:t')
    t.text = '1'
    run_text.append(t)
    r_element.append(run_text)
    
    # Конец поля
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r_element.append(fldChar_end)
    
    # Форматирование
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)

def parse_markdown(md_file):
    """Парсит markdown файл и возвращает список элементов"""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    elements = []
    current_paragraph = []
    in_list = False
    is_title_page = True  # Первые строки - титульный лист
    
    for line in lines:
        line = line.rstrip('\n\r')
        
        # Placeholder для оглавления
        if line.strip() == '<!-- TOC -->':
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('toc_placeholder',))
            continue
        
        # Разрыв страницы (---)
        if line.strip() == '---':
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('page_break',))
            is_title_page = False  # После первого разрыва - не титульный лист
            continue
        
        # Заголовки (# ## ###)
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            
            # Если это первый заголовок на титульном листе, обрабатываем как обычный текст
            if is_title_page and level == 1 and text == 'home-sentinel':
                elements.append(('title', text))
            else:
                elements.append(('heading', level, text))
                is_title_page = False
            continue
        
        # Списки (- или *)
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            in_list = True
            indent = list_match.group(1)
            text = list_match.group(2)
            # Проверяем вложенность (отступы: 4 пробела = 1 уровень)
            indent_level = len(indent) // 4
            elements.append(('list_item', text, indent_level))
            continue
        
        # Нумерованные списки (1. 2. и т.д.)
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            in_list = True
            indent = numbered_match.group(1)
            text = numbered_match.group(2)
            indent_level = len(indent) // 4
            elements.append(('numbered_item', text, indent_level))
            continue
        
        # Пустая строка
        if not line.strip():
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            in_list = False
            continue
        
        # Обычный текст
        if in_list:
            # Продолжение списка (многострочный элемент)
            if current_paragraph:
                current_paragraph.append(line)
            else:
                current_paragraph = [line]
        else:
            current_paragraph.append(line)
    
    # Добавляем последний параграф
    if current_paragraph:
        elements.append(('paragraph', '\n'.join(current_paragraph)))
    
    return elements

def generate_doc_from_markdown(md_file, output_file):
    """Генерирует DOCX документ из markdown файла"""
    print(f"📄 Чтение {md_file}...")
    elements = parse_markdown(md_file)
    
    # Первый проход: собираем все заголовки для оглавления
    print("📋 Сбор заголовков для оглавления...")
    headings_list = []
    for element in elements:
        if element[0] == 'heading':
            level = element[1]
            text = element[2]
            headings_list.append((level, text))
    print(f"   Найдено заголовков: {len(headings_list)}")
    
    # Создаем документ
    doc = Document()
    
    # Настройка полей: слева 2 см, справа 1 см, сверху и снизу 2 см
    # И добавление нумерации страниц
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        # Добавляем нумерацию страниц в нижний колонтитул
        add_page_number_footer(section)
    
    print("📝 Генерация документа...")
    
    for i, element in enumerate(elements):
        elem_type = element[0]
        
        if elem_type == 'page_break':
            doc.add_page_break()
            print(f"  Разрыв страницы ({i+1}/{len(elements)})")
        
        elif elem_type == 'toc_placeholder':
            print(f"  🔍 Создание оглавления ({i+1}/{len(elements)})...")
            
            # Добавляем заголовок "Содержание"
            add_formatted_heading(doc, 'Содержание', 1)
            
            # Пробуем создать TOC поле (работает в Word и может работать в LibreOffice)
            print("    - Создание TOC поля...")
            para = doc.add_paragraph()
            
            # Создаем TOC поле в правильной структуре (каждый элемент в отдельном run)
            run1 = para.add_run()
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            run1._element.append(fldChar_begin)
            
            run2 = para.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
            run2._element.append(instrText)
            
            run3 = para.add_run()
            fldChar_separate = OxmlElement('w:fldChar')
            fldChar_separate.set(qn('w:fldCharType'), 'separate')
            run3._element.append(fldChar_separate)
            
            run4 = para.add_run('Нажмите F9 или правой кнопкой -> Обновить поле')
            
            run5 = para.add_run()
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run5._element.append(fldChar_end)
            
            print(f"  ✅ TOC поле создано")
            print("  ⚠️  В Word: нажмите F9 для обновления")
            print("  ⚠️  В LibreOffice: правой кнопкой на оглавлении -> Обновить оглавление/указатель")
        
        elif elem_type == 'title':
            # Титульный лист - заголовок
            text = element[1]
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(text)
            title_run.font.name = FONT_NAME
            title_run.font.size = Pt(18)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"  Титульный лист: {text} ({i+1}/{len(elements)})")
        
        elif elem_type == 'heading':
            level = element[1]
            text = element[2]
            add_formatted_heading(doc, text, level)
            print(f"  Заголовок {level}: {text[:50]}... ({i+1}/{len(elements)})")
        
        elif elem_type == 'paragraph':
            text = element[1].strip()
            if not text:
                continue
            
            # Специальная обработка для титульного листа (первые несколько элементов)
            # Проверяем, не является ли это частью титульного листа
            if i < 5 and (text == 'Home Assistant AI Stack с поддержкой GPU' or 
                        text == 'Интеллектуальная система мониторинга и распознавания для умного дома'):
                if text == 'Home Assistant AI Stack с поддержкой GPU':
                    # Подзаголовок титульного листа
                    subtitle_para = doc.add_paragraph()
                    subtitle_run = subtitle_para.add_run(text)
                    subtitle_run.font.name = FONT_NAME
                    subtitle_run.font.size = Pt(14)
                    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                elif text == 'Интеллектуальная система мониторинга и распознавания для умного дома':
                    # Описание титульного листа
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(text)
                    desc_run.font.name = FONT_NAME
                    desc_run.font.size = Pt(FONT_SIZE)
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
            
            # Обычный параграф
            add_formatted_paragraph(doc, text)
        
        elif elem_type == 'list_item':
            text = element[1]
            indent_level = element[2]
            style = 'List Bullet 2' if indent_level > 0 else 'List Bullet'
            para = doc.add_paragraph(text, style=style)
            set_paragraph_formatting(para, alignment='left')
        
        elif elem_type == 'numbered_item':
            text = element[1]
            indent_level = element[2] if len(element) > 2 else 0
            # Для вложенных нумерованных списков используем List Number 2
            style = 'List Number 2' if indent_level > 0 else 'List Number'
            para = doc.add_paragraph(text, style=style)
            set_paragraph_formatting(para, alignment='left')
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"\n✅ Документ создан: {output_file}")
    print(f"   Размер: {os.path.getsize(output_file) / 1024:.1f} KB")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Генерация документа проекта home-sentinel из project.md'
    )
    parser.add_argument(
        '--publish',
        action='store_true',
        help='Подготовить HTML для публикации в GitHub Pages (только для HTML генераторов)'
    )
    parser.add_argument(
        '--release',
        action='store_true',
        help='Создать GitHub Release с сгенерированным документом (DOCX/PPTX)'
    )
    parser.add_argument(
        '--version',
        type=str,
        help='Версия для GitHub Release (по умолчанию: timestamp)'
    )
    parser.add_argument(
        '--input',
        type=str,
        help='Путь к исходному markdown файлу (по умолчанию: docs/project.md)'
    )
    parser.add_argument(
        '--output',
        type=str,
        help='Путь к выходному файлу (по умолчанию: используется BUILD_DIR и шаблон из .env)'
    )
    
    args = parser.parse_args()
    
    # Загружаем конфигурацию
    env_vars = load_env_file()
    
    # Определяем пути
    project_root = get_project_root()
    
    # Путь к исходному файлу
    if args.input:
        md_file = args.input
    else:
        md_file = project_root / env_vars.get('DOCS_PROJECT_MD', 'docs/project.md')
    
    if not os.path.exists(md_file):
        print(f"❌ Файл не найден: {md_file}")
        sys.exit(1)
    
    # Путь к выходному файлу
    if args.output:
        output_file = args.output
    else:
        # Генерируем имя файла по шаблону
        filename, ext = generate_output_filename(str(md_file), env_vars=env_vars, output_ext='.docx')
        build_dir = get_build_dir(env_vars)
        output_file = build_dir / f"{filename}{ext}"
    
    # Создаем директорию для выходного файла если нужно
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # Генерируем документ
    generate_doc_from_markdown(str(md_file), str(output_file))
    
    # Создаем GitHub Release если нужно (для DOCX/PPTX)
    if args.release:
        create_github_release([str(output_file)], env_vars, args.version)
